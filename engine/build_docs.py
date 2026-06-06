"""
A4 — Pre-bite and post-bite document builder (Person A).

Public API (do not change signatures without team sync):
    build_bites(spec, out_dir) -> tuple[str, str]
"""
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt


def _markdown_to_docx(doc: Document, markdown: str) -> None:
    """Minimal markdown → docx: handles ##/# headings and bullet lines starting with -/•."""
    for line in markdown.splitlines():
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith(("- ", "• ")):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
            doc.add_paragraph(line[3:].strip(), style="List Number")
        else:
            doc.add_paragraph(line)


def build_bites(spec: dict, out_dir: str) -> tuple[str, str]:
    """
    Writes pre-bite and post-bite as .docx files.
    Returns (prebite_path, postbite_path).
    """
    Path(out_dir).mkdir(parents=True, exist_ok=True)
    topic = spec["meta"].get("topic", "Training")

    prebite_path  = str(Path(out_dir) / "prebite.docx")
    postbite_path = str(Path(out_dir) / "postbite.docx")

    for content, path, label in [
        (spec["prebite"],  prebite_path,  "Pre-Session Preparation"),
        (spec["postbite"], postbite_path, "Post-Session Follow-Up"),
    ]:
        doc = Document()
        doc.add_heading(f"{label}: {topic}", level=1)
        _markdown_to_docx(doc, content)
        doc.save(path)

    return prebite_path, postbite_path


if __name__ == "__main__":
    import sys
    spec_path = sys.argv[1] if len(sys.argv) > 1 else "schema/sample_spec.json"
    out_dir   = sys.argv[2] if len(sys.argv) > 2 else "output"

    spec = json.loads(Path(spec_path).read_text())
    pre, post = build_bites(spec, out_dir)
    print(f"Pre-bite:  {pre}")
    print(f"Post-bite: {post}")
